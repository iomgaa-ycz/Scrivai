"""DocxRenderer — docxtpl template rendering.

Constraints (docxtpl limitations):
1. Templates must be hand-crafted in Word/LibreOffice (jinja tags must be in a single <w:r>).
2. Nested {% for %} inside a single cell is not supported; use jinja filters to flatten.
3. Avoid tables-within-tables.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docxtpl import DocxTemplate

# Regex matching the var name inside a docxtpl placeholder {{ var }}
_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z_][a-zA-Z0-9_]*)\s*\}\}")


class DocxRenderer:
    """docx template renderer based on docxtpl."""

    def __init__(self, template_path: str | Path) -> None:
        self._template_path = Path(template_path)
        if not self._template_path.is_file():
            raise FileNotFoundError(f"docx template not found: {self._template_path}")
        # Load once to validate that the file can be parsed by docxtpl
        self._template = DocxTemplate(str(self._template_path))

    @property
    def template_path(self) -> Path:
        return self._template_path

    def list_placeholders(self) -> list[str]:
        """Scan the template for all {{ var }} placeholders and return a sorted, deduplicated list."""
        # Plain-text extraction from the internal docx XML
        names: set[str] = set()
        # Re-open to get raw text (does not depend on docxtpl's jinja env)
        from docx import Document

        doc = Document(str(self._template_path))
        for para in doc.paragraphs:
            names.update(_PLACEHOLDER_RE.findall(para.text))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        names.update(_PLACEHOLDER_RE.findall(para.text))
        return sorted(names)

    def render(self, context: dict[str, Any], output_path: str | Path) -> Path:
        """Render the template to output_path; deletes any partial output on failure.

        The template is reloaded on each render call (DocxTemplate.render is stateful).
        """
        out = Path(output_path)
        # Ensure the parent directory exists; raise if it does not (test expectation)
        if not out.parent.is_dir():
            raise IOError(f"output directory does not exist: {out.parent}")

        # Re-open the template (DocxTemplate rendering is stateful)
        tpl = DocxTemplate(str(self._template_path))
        try:
            tpl.render(context)
            tpl.save(str(out))
        except Exception:
            # Any exception: remove the partial output then re-raise
            if out.exists():
                try:
                    out.unlink()
                except OSError:
                    pass
            raise
        return out
