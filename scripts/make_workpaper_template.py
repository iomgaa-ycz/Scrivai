"""一次性构造 tests/fixtures/m1_e2e/workpaper_template.docx(docxtpl 3 占位符)。

3 个占位符均为简单标识符(不含 . / Jinja 复杂标签),与 DocxRenderer 的
list_placeholders 正则兼容。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

FIXTURE = (
    Path(__file__).resolve().parents[1]
    / "tests"
    / "fixtures"
    / "m1_e2e"
    / "workpaper_template.docx"
)


def build() -> None:
    FIXTURE.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("变电站技术监督审核工作底稿", level=1)
    doc.add_paragraph("项目名称:{{ project_name }}")
    doc.add_paragraph("报告日期:{{ report_date }}")
    doc.add_heading("审核结论", level=2)
    doc.add_paragraph("{{ audit_summary }}")
    doc.save(FIXTURE)
    print(f"built: {FIXTURE}")


if __name__ == "__main__":
    build()
