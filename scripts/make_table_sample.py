"""一次性构造 tests/fixtures/io_samples/table_sample.docx。

内容:标题 + 前文段 + 2 列 3 行表(含表头)+ 后文段。
用于 test_docx_to_markdown_preserves_table 验证 pandoc 保留表格结构。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

FIXTURE = (
    Path(__file__).resolve().parents[1] / "tests" / "fixtures" / "io_samples" / "table_sample.docx"
)


def build() -> None:
    FIXTURE.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Substation Inspection Report", level=1)
    doc.add_paragraph("Pre-table narrative paragraph.")

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Checkpoint"
    table.rows[0].cells[1].text = "Verdict"
    table.rows[1].cells[0].text = "Insulation resistance"
    table.rows[1].cells[1].text = "Pass"
    table.rows[2].cells[0].text = "Grounding continuity"
    table.rows[2].cells[1].text = "Fail"

    doc.add_paragraph("Post-table conclusion.")
    doc.save(FIXTURE)
    print(f"built {FIXTURE} ({FIXTURE.stat().st_size} B)")


if __name__ == "__main__":
    build()
