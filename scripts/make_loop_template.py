"""一次性构造 tests/fixtures/io_samples/loop_template.docx。

docxtpl 的 {% for %} 标签必须在单一 <w:r> 内,python-docx 生成时可能拆 run。
规避:先写占位串(__TAG_FOR__ / __TAG_ITEM_NAME__ / __TAG_ENDFOR__),
保存后 unzip docx,字符串替换 word/document.xml,再重新打包。

目标模板:
  Report: {{ project_name }}

  {% for item in items %}
  {{ item.name }}
  {% endfor %}
"""

from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

FIXTURE = (
    Path(__file__).resolve().parents[1] / "tests" / "fixtures" / "io_samples" / "loop_template.docx"
)

PLACEHOLDER_FOR = "__TAG_FOR__"
PLACEHOLDER_ITEM = "__TAG_ITEM_NAME__"
PLACEHOLDER_ENDFOR = "__TAG_ENDFOR__"

REPLACEMENTS: dict[str, str] = {
    PLACEHOLDER_FOR: "{% for item in items %}",
    PLACEHOLDER_ITEM: "{{ item.name }}",
    PLACEHOLDER_ENDFOR: "{% endfor %}",
}


def build() -> None:
    FIXTURE.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("Report: {{ project_name }}", level=1)
    doc.add_paragraph(PLACEHOLDER_FOR)
    doc.add_paragraph(PLACEHOLDER_ITEM)
    doc.add_paragraph(PLACEHOLDER_ENDFOR)

    tmp = FIXTURE.with_suffix(".tmp.docx")
    doc.save(tmp)

    # XML 后处理:解压 → 替换 → 重新打包
    with zipfile.ZipFile(tmp, "r") as zin:
        names = zin.namelist()
        blobs = {n: zin.read(n) for n in names}

    xml = blobs["word/document.xml"].decode("utf-8")
    for src, dst in REPLACEMENTS.items():
        if src not in xml:
            raise RuntimeError(f"占位符 {src!r} 未出现在 document.xml — python-docx 可能已拆 run")
        xml = xml.replace(src, dst)
    blobs["word/document.xml"] = xml.encode("utf-8")

    with zipfile.ZipFile(FIXTURE, "w", zipfile.ZIP_DEFLATED) as zout:
        for n in names:
            zout.writestr(n, blobs[n])

    tmp.unlink()
    print(f"built {FIXTURE} ({FIXTURE.stat().st_size} B)")


if __name__ == "__main__":
    build()
