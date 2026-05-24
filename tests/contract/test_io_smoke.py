"""Contract tests for scrivai.io — unified to_markdown() + DocxRenderer.

Uses real_data/ files for OCR pipeline smoke tests.
"""

from __future__ import annotations

import shutil
from pathlib import Path

import pytest

REAL_DATA = Path(__file__).resolve().parents[2] / "real_data"
FIXTURES_DIR = Path(__file__).resolve().parents[1] / "fixtures" / "io_samples"

SAMPLE_DOC = REAL_DATA / "2025年政府采购领域“四类”违法违规行为专项整治工作指引.doc"
SAMPLE_DOCX = (
    REAL_DATA
    / "从化区中医医院手术室设备及附件、病房护理及医院设备采购"
    / "从化区中医医院手术室设备及附件、病房护理及医院设备采购.docx"
)
SAMPLE_PDF = (
    REAL_DATA
    / "从化区中医医院手术室设备及附件、病房护理及医院设备采购"
    / "3、从化区中医医院手术室设备及附件、病房护理及医院设备采购"
    / "从化区中医医院手术室设备及附件、病房护理及医院设备采购招标文件（2024040902）.pdf.pdf"
)


def _monkeyocr_reachable(base_url: str = "http://100.81.95.44:7861") -> bool:
    import requests

    try:
        session = requests.Session()
        session.trust_env = False
        session.get(base_url, timeout=2)
        return True
    except requests.exceptions.RequestException:
        return False


def _has_soffice() -> bool:
    return bool(shutil.which("libreoffice") or shutil.which("soffice"))


# ─── to_markdown: OCR 主路径 ──────────────────────────────────


@pytest.mark.skipif(not _monkeyocr_reachable(), reason="MonkeyOCR 不可达")
@pytest.mark.skipif(not _has_soffice(), reason="需要 libreoffice 二进制")
def test_to_markdown_doc() -> None:
    """真实 .doc 文件经 LibreOffice → PDF → MonkeyOCR 输出 Markdown。"""
    from scrivai.io import to_markdown

    if not SAMPLE_DOC.is_file():
        pytest.skip(f"测试文件不存在: {SAMPLE_DOC}")

    md = to_markdown(SAMPLE_DOC, timeout=300)
    assert isinstance(md, str)
    assert len(md) > 100
    assert "政府采购" in md


@pytest.mark.skipif(not _monkeyocr_reachable(), reason="MonkeyOCR 不可达")
@pytest.mark.skipif(not _has_soffice(), reason="需要 libreoffice 二进制")
def test_to_markdown_docx() -> None:
    """真实 .docx 文件经 LibreOffice → PDF → MonkeyOCR 输出 Markdown。"""
    from scrivai.io import to_markdown

    if not SAMPLE_DOCX.is_file():
        pytest.skip(f"测试文件不存在: {SAMPLE_DOCX}")

    md = to_markdown(SAMPLE_DOCX, timeout=300)
    assert isinstance(md, str)
    assert len(md) > 100


@pytest.mark.skipif(not _monkeyocr_reachable(), reason="MonkeyOCR 不可达")
def test_to_markdown_pdf() -> None:
    """真实 PDF 直接走 MonkeyOCR。"""
    from scrivai.io import to_markdown

    if not SAMPLE_PDF.is_file():
        pytest.skip(f"测试文件不存在: {SAMPLE_PDF}")

    md = to_markdown(SAMPLE_PDF, timeout=300)
    assert isinstance(md, str)
    assert len(md) > 100


# ─── to_markdown: fallback 路径 ───────────────────────────────


@pytest.mark.skipif(not shutil.which("pandoc"), reason="需要 pandoc 二进制")
def test_to_markdown_fallback_docx(tmp_path: Path) -> None:
    """MonkeyOCR 不可达时 .docx 降级到 pandoc。"""
    from docx import Document

    from scrivai.io import to_markdown

    doc = Document()
    doc.add_heading("Fallback Test", level=1)
    doc.add_paragraph("Hello fallback world.")
    fixture = tmp_path / "fallback.docx"
    doc.save(fixture)

    md = to_markdown(
        fixture,
        ocr_base_url="http://127.0.0.1:1",
        timeout=2,
        fallback=True,
    )
    assert isinstance(md, str)
    assert "Hello fallback world" in md


# ─── to_markdown: 错误路径 ────────────────────────────────────


def test_to_markdown_unsupported(tmp_path: Path) -> None:
    """不支持的格式 → IOError。"""
    from scrivai.io import to_markdown

    fake = tmp_path / "data.xls"
    fake.write_bytes(b"\x00\x01\x02")

    with pytest.raises(IOError, match="Unsupported format"):
        to_markdown(fake)


def test_to_markdown_file_not_found(tmp_path: Path) -> None:
    """文件不存在 → IOError。"""
    from scrivai.io import to_markdown

    with pytest.raises(IOError, match="File not found"):
        to_markdown(tmp_path / "no-such-file.pdf")


def test_to_markdown_invalid_backend(tmp_path: Path) -> None:
    """非法后端名 → ValueError。"""
    from scrivai.io import to_markdown

    fake = tmp_path / "data.pdf"
    fake.write_bytes(b"%PDF-1.4 fake")
    with pytest.raises(ValueError, match="Unknown OCR backend"):
        to_markdown(fake, ocr_backend="nonexistent")


# ─── DocxRenderer ─────────────────────────────────────────────


@pytest.fixture
def sample_template(tmp_path: Path) -> Path:
    """用 python-docx 造一个含 docxtpl 占位符的 .docx 模板。"""
    from docx import Document

    doc = Document()
    doc.add_heading("Project: {{ project_name }}", level=1)
    doc.add_paragraph("Author: {{ author }}")
    doc.add_paragraph("Body: {{ body }}")
    out = tmp_path / "template.docx"
    doc.save(out)
    return out


def test_docx_renderer_list_placeholders(sample_template: Path) -> None:
    """list_placeholders 返回去重排序的占位符名列表。"""
    from scrivai.io import DocxRenderer

    renderer = DocxRenderer(sample_template)
    names = renderer.list_placeholders()
    assert names == sorted(set(names))
    assert {"project_name", "author", "body"}.issubset(set(names))


def test_docx_renderer_render(sample_template: Path, tmp_path: Path) -> None:
    """render 写出 docx,文件存在且非空。"""
    from scrivai.io import DocxRenderer

    out = tmp_path / "rendered.docx"
    renderer = DocxRenderer(sample_template)
    result = renderer.render(
        context={"project_name": "X变电站", "author": "yu", "body": "hello"},
        output_path=out,
    )
    assert result == out
    assert out.is_file()
    assert out.stat().st_size > 0


def test_docx_renderer_template_not_found(tmp_path: Path) -> None:
    """模板不存在 → FileNotFoundError。"""
    from scrivai.io import DocxRenderer

    with pytest.raises(FileNotFoundError):
        DocxRenderer(tmp_path / "no-such.docx")


def test_docx_renderer_render_failure_no_halfproduct(sample_template: Path, tmp_path: Path) -> None:
    """渲染异常时不留半成品文件。"""
    from scrivai.io import DocxRenderer

    bad_out = tmp_path / "no-such-dir" / "x.docx"
    renderer = DocxRenderer(sample_template)

    with pytest.raises((IOError, OSError, FileNotFoundError)):
        renderer.render(context={"project_name": "p"}, output_path=bad_out)

    assert not bad_out.exists()


# ─── DocxRenderer: fixture edge cases ─────────────────────────


def test_docx_renderer_list_placeholders_in_loop() -> None:
    """loop_template.docx 含简单占位 + 复杂占位;只断言简单标识符出现。"""
    from scrivai.io import DocxRenderer

    fixture = FIXTURES_DIR / "loop_template.docx"
    assert fixture.is_file(), f"fixture 不存在:{fixture}"

    renderer = DocxRenderer(fixture)
    names = renderer.list_placeholders()
    assert "project_name" in names, f"project_name 必在;actual={names}"


def test_docx_renderer_render_loop(tmp_path: Path) -> None:
    """loop_template.docx 渲染 3 item 后,每 item 应在输出中各出现一次。"""
    from docx import Document as _Doc

    from scrivai.io import DocxRenderer

    fixture = FIXTURES_DIR / "loop_template.docx"
    out = tmp_path / "loop_rendered.docx"

    renderer = DocxRenderer(fixture)
    ctx: dict = {
        "project_name": "Sub-7",
        "items": [{"name": "alpha"}, {"name": "beta"}, {"name": "gamma"}],
    }
    result = renderer.render(context=ctx, output_path=out)
    assert result == out
    assert out.is_file()
    assert out.stat().st_size > 0

    rendered = _Doc(str(out))
    text = "\n".join(p.text for p in rendered.paragraphs)
    assert "Sub-7" in text, f"project_name 未渲染;text:\n{text}"
    for name in ("alpha", "beta", "gamma"):
        assert name in text, f"item {name!r} 未渲染;text:\n{text}"
