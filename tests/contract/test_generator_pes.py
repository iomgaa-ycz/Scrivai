"""GeneratorPES contract tests (M1.5a T1.6)."""

from __future__ import annotations

import json
import os
import zipfile
from datetime import datetime, timezone
from pathlib import Path

import pytest
from docx import Document
from pydantic import BaseModel

from scrivai.agents.generator import GeneratorPES
from scrivai.models.pes import (
    ModelConfig,
    PESConfig,
    PESRun,
    PhaseResult,
)
from scrivai.models.workspace import WorkspaceHandle, WorkspaceSnapshot
from scrivai.pes.config import load_pes_config

# ──────────────── Fixtures ────────────────


class _Section(BaseModel):
    placeholder: str
    content: str


class _GenOutput(BaseModel):
    context: dict[str, str]
    sections: list[_Section]


@pytest.fixture
def workspace(tmp_path: Path) -> WorkspaceHandle:
    project = tmp_path / "project"
    project.mkdir()
    ws = tmp_path / "ws"
    for sub in ("working", "data", "output", "logs"):
        (ws / sub).mkdir(parents=True)
    return WorkspaceHandle(
        run_id="r1",
        root_dir=ws,
        working_dir=ws / "working",
        data_dir=ws / "data",
        output_dir=ws / "output",
        logs_dir=ws / "logs",
        snapshot=WorkspaceSnapshot(
            run_id="r1",
            project_root=project,
            snapshot_at=datetime.now(timezone.utc),
        ),
    )


@pytest.fixture
def config() -> PESConfig:
    return load_pes_config(Path("scrivai/agents/generator.yaml"))


@pytest.fixture
def template_path(tmp_path: Path) -> Path:
    """生成最小 docxtpl 模板:含 {{ project_name }} 和 {{ report_date }}。"""
    tpl_path = tmp_path / "tmpl.docx"
    doc = Document()
    doc.add_paragraph("Project: {{ project_name }}")
    doc.add_paragraph("Date: {{ report_date }}")
    doc.save(tpl_path)
    return tpl_path


def _phase_result(phase: str = "summarize") -> PhaseResult:
    return PhaseResult(phase=phase, attempt_no=0, started_at=datetime.now(timezone.utc))


def _run(pes: GeneratorPES) -> PESRun:
    return PESRun(
        run_id="r1",
        pes_name="generator",
        task_prompt="test",
        model_name="mock",
        started_at=datetime.now(timezone.utc),
    )


# ──────────────── build_execution_context ────────────────


async def test_build_execution_context_plan_parses_placeholders(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"template_path": template_path, "context_schema": _GenOutput},
    )
    ctx = await pes.build_execution_context("plan", _run(pes))
    assert "placeholders" in ctx
    assert set(ctx["placeholders"]) == {"project_name", "report_date"}


async def test_build_execution_context_non_plan_empty(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"template_path": template_path, "context_schema": _GenOutput},
    )
    ctx_execute = await pes.build_execution_context("execute", _run(pes))
    ctx_summarize = await pes.build_execution_context("summarize", _run(pes))
    assert ctx_execute == {}
    assert ctx_summarize == {}


async def test_build_execution_context_requires_template_path(workspace, config):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"context_schema": _GenOutput},
    )
    with pytest.raises(ValueError, match="template_path"):
        await pes.build_execution_context("plan", _run(pes))


# ──────────────── postprocess_phase_result ────────────────


async def test_postprocess_happy_path_auto_render_false(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={
            "template_path": template_path,
            "context_schema": _GenOutput,
        },
    )
    (workspace.working_dir / "output.json").write_text(
        json.dumps(
            {
                "context": {"project_name": "XX变电站", "report_date": "2026-04-17"},
                "sections": [
                    {"placeholder": "project_name", "content": "XX变电站"},
                    {"placeholder": "report_date", "content": "2026-04-17"},
                ],
            }
        ),
        encoding="utf-8",
    )
    run = _run(pes)
    await pes.postprocess_phase_result("summarize", _phase_result(), run)
    assert run.final_output is not None
    assert not (workspace.output_dir / "final.docx").exists()


async def test_postprocess_auto_render_true_produces_docx(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={
            "template_path": template_path,
            "context_schema": _GenOutput,
            "auto_render": True,
        },
    )
    (workspace.working_dir / "output.json").write_text(
        json.dumps(
            {
                "context": {"project_name": "XX变电站", "report_date": "2026-04-17"},
                "sections": [
                    {"placeholder": "project_name", "content": "XX变电站"},
                    {"placeholder": "report_date", "content": "2026-04-17"},
                ],
            }
        ),
        encoding="utf-8",
    )
    await pes.postprocess_phase_result("summarize", _phase_result(), _run(pes))
    docx_path = workspace.output_dir / "final.docx"
    assert docx_path.exists()
    with zipfile.ZipFile(docx_path) as zf:
        content_xml = zf.read("word/document.xml").decode("utf-8")
    assert "XX变电站" in content_xml
    assert "2026-04-17" in content_xml


async def test_postprocess_missing_context_schema(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"template_path": template_path},
    )
    (workspace.working_dir / "output.json").write_text(
        json.dumps({"context": {}, "sections": []}), encoding="utf-8"
    )
    with pytest.raises(ValueError, match="context_schema"):
        await pes.postprocess_phase_result("summarize", _phase_result(), _run(pes))


async def test_postprocess_missing_template_path(workspace, config):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"context_schema": _GenOutput},
    )
    (workspace.working_dir / "output.json").write_text(
        json.dumps({"context": {}, "sections": []}), encoding="utf-8"
    )
    with pytest.raises(ValueError, match="template_path"):
        await pes.postprocess_phase_result("summarize", _phase_result(), _run(pes))


# ──────────────── validate_phase_outputs ────────────────


async def test_validate_plan_coverage_ok(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"template_path": template_path, "context_schema": _GenOutput},
    )
    (workspace.working_dir / "plan.md").write_text("plan", encoding="utf-8")
    (workspace.working_dir / "plan.json").write_text(
        json.dumps(
            {
                "fills": [
                    {"placeholder": "project_name", "source": "a"},
                    {"placeholder": "report_date", "source": "b"},
                ]
            }
        ),
        encoding="utf-8",
    )
    await pes.validate_phase_outputs(
        "plan", config.phases["plan"], _phase_result("plan"), _run(pes)
    )


async def test_validate_plan_placeholder_missing(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"template_path": template_path, "context_schema": _GenOutput},
    )
    (workspace.working_dir / "plan.md").write_text("plan", encoding="utf-8")
    (workspace.working_dir / "plan.json").write_text(
        json.dumps({"fills": [{"placeholder": "project_name", "source": "a"}]}),
        encoding="utf-8",
    )
    with pytest.raises(ValueError, match="report_date"):
        await pes.validate_phase_outputs(
            "plan", config.phases["plan"], _phase_result("plan"), _run(pes)
        )


async def test_validate_execute_coverage_ok(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"template_path": template_path, "context_schema": _GenOutput},
    )
    findings = workspace.working_dir / "findings"
    findings.mkdir()
    (findings / "project_name.json").write_text("{}", encoding="utf-8")
    (findings / "report_date.json").write_text("{}", encoding="utf-8")
    await pes.validate_phase_outputs(
        "execute", config.phases["execute"], _phase_result("execute"), _run(pes)
    )


async def test_validate_execute_placeholder_missing(workspace, config, template_path):
    pes = GeneratorPES(
        config=config,
        model=ModelConfig(model="mock"),
        workspace=workspace,
        runtime_context={"template_path": template_path, "context_schema": _GenOutput},
    )
    findings = workspace.working_dir / "findings"
    findings.mkdir()
    (findings / "project_name.json").write_text("{}", encoding="utf-8")
    with pytest.raises(ValueError, match="report_date"):
        await pes.validate_phase_outputs(
            "execute", config.phases["execute"], _phase_result("execute"), _run(pes)
        )


# ──────────────── Real GLM smoke ────────────────


@pytest.mark.skipif(
    not os.getenv("ANTHROPIC_AUTH_TOKEN"),
    reason="real SDK smoke; requires ANTHROPIC_AUTH_TOKEN in .env",
)
async def test_generator_smoke_with_real_glm(workspace, config, template_path):
    """GeneratorPES 真跑 GLM-5.1:2 占位符模板 + 短素材 → final.docx 渲染。"""
    (workspace.data_dir / "source.md").write_text(
        "# Project\nXX变电站扩建项目\n\n# Date\n2026-04-17\n",
        encoding="utf-8",
    )

    model = ModelConfig(
        model=os.getenv("SCRIVAI_DEFAULT_MODEL") or "glm-5.1",
        provider=os.getenv("SCRIVAI_DEFAULT_PROVIDER") or "glm",
    )
    pes = GeneratorPES(
        config=config,
        model=model,
        workspace=workspace,
        runtime_context={
            "template_path": template_path,
            "context_schema": _GenOutput,
            "auto_render": True,
        },
    )
    run = await pes.run(
        "Read data/source.md. Fill the 'project_name' placeholder with the "
        "text under '# Project' and 'report_date' with the text under '# Date'. "
        "Both values are short single-line strings."
    )
    assert run.status == "completed", f"run failed: {run.error_type} {run.error}"
    assert (workspace.output_dir / "final.docx").exists()
