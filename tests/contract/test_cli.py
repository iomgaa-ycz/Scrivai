"""M0.75 T0.13 contract tests for scrivai-cli。

参考 docs/superpowers/specs/2026-04-16-scrivai-m0.75-design.md §5。
"""

from __future__ import annotations

import json
import os
import subprocess
import sys
from pathlib import Path

import pytest

CLI_CMD = [sys.executable, "-m", "scrivai.cli"]


def _run(args: list[str], env: dict | None = None) -> tuple[int, str, str]:
    env_full = {**os.environ, **(env or {})}
    proc = subprocess.run(CLI_CMD + args, capture_output=True, text=True, env=env_full)
    return proc.returncode, proc.stdout, proc.stderr


def _parse_error_json(stderr: str) -> dict:
    """从 stderr 提取最后一条 JSON 错误对象,容忍前面的 warning/log。"""
    import json as _json

    for line in reversed(stderr.strip().splitlines()):
        line = line.strip()
        if line.startswith("{") and line.endswith("}"):
            try:
                return _json.loads(line)
            except _json.JSONDecodeError:
                continue
    raise AssertionError(f"no JSON error object in stderr:\n{stderr}")


# ── library group ──


@pytest.fixture
def populated_qmd(tmp_path: Path):
    """造一个带几条数据的 qmd db。"""
    from scrivai.knowledge import build_libraries, build_qmd_client_from_config

    db = tmp_path / "test.db"
    client = build_qmd_client_from_config(db)
    rules, cases, _ = build_libraries(client)
    rules.add(entry_id="r-1", markdown="围标串标的认定标准", metadata={"law": "32"})
    rules.add(entry_id="r-2", markdown="供应商资格", metadata={})
    cases.add(entry_id="c-1", markdown="案例 A", metadata={})
    return db


def test_library_list_json_shape(populated_qmd: Path) -> None:
    code, out, err = _run(["library", "list", "--type", "rules", "--db-path", str(populated_qmd)])
    assert code == 0, err
    payload = json.loads(out)
    assert "entry_ids" in payload
    assert set(payload["entry_ids"]) == {"r-1", "r-2"}


def test_library_get_json_shape(populated_qmd: Path) -> None:
    code, out, err = _run(
        [
            "library",
            "get",
            "--type",
            "rules",
            "--db-path",
            str(populated_qmd),
            "--entry-id",
            "r-1",
        ]
    )
    assert code == 0, err
    payload = json.loads(out)
    assert payload["entry_id"] == "r-1"
    assert "围标串标" in payload["markdown"]
    assert payload["metadata"] == {"law": "32"}


def test_library_get_missing_returns_error(populated_qmd: Path) -> None:
    code, out, err = _run(
        [
            "library",
            "get",
            "--type",
            "rules",
            "--db-path",
            str(populated_qmd),
            "--entry-id",
            "no-such",
        ]
    )
    assert code == 1
    payload = _parse_error_json(err)
    assert "error" in payload
    assert "no-such" in payload["error"]


def test_library_missing_env_gives_error_json(tmp_path: Path) -> None:
    """没有 --db-path 也没有 QMD_DB_PATH env → stderr JSON + exit 1。"""
    env_clean = {k: v for k, v in os.environ.items() if k != "QMD_DB_PATH"}
    proc = subprocess.run(
        CLI_CMD + ["library", "list", "--type", "rules"],
        capture_output=True,
        text=True,
        env=env_clean,
    )
    assert proc.returncode == 1
    payload = _parse_error_json(proc.stderr)
    assert "error" in payload
    assert "QMD_DB_PATH" in payload["error"]


def test_library_search_json_shape(populated_qmd: Path) -> None:
    code, out, err = _run(
        [
            "library",
            "search",
            "--type",
            "rules",
            "--db-path",
            str(populated_qmd),
            "--query",
            "围标",
            "--top-k",
            "2",
        ]
    )
    assert code == 0, err
    payload = json.loads(out)
    assert "hits" in payload
    assert isinstance(payload["hits"], list)


# ── io group ──


def test_cli_docx2md_writes_output(tmp_path: Path) -> None:
    """CLI docx2md 写出 markdown 文件。"""
    import shutil

    if not shutil.which("pandoc"):
        pytest.skip("需要 pandoc")

    from docx import Document

    doc = Document()
    doc.add_heading("CLI Test", level=1)
    doc.add_paragraph("Hello CLI.")
    src = tmp_path / "in.docx"
    doc.save(src)
    out = tmp_path / "out.md"

    code, stdout, err = _run(["io", "docx2md", "--input", str(src), "--output", str(out)])
    assert code == 0, err
    payload = json.loads(stdout)
    assert payload["output"] == str(out)
    assert "Hello CLI" in out.read_text(encoding="utf-8")


def test_cli_render_produces_docx(tmp_path: Path) -> None:
    """CLI render 用 docxtpl 模板 + JSON context → 写 docx。"""
    from docx import Document

    tpl = Document()
    tpl.add_paragraph("Hello {{ name }}")
    tpl_path = tmp_path / "tpl.docx"
    tpl.save(tpl_path)

    ctx_path = tmp_path / "ctx.json"
    ctx_path.write_text(json.dumps({"name": "scrivai"}), encoding="utf-8")
    out_path = tmp_path / "out.docx"

    code, stdout, err = _run(
        [
            "io",
            "render",
            "--template",
            str(tpl_path),
            "--context-json",
            str(ctx_path),
            "--output",
            str(out_path),
        ]
    )
    assert code == 0, err
    payload = json.loads(stdout)
    assert payload["output"] == str(out_path)
    assert out_path.is_file()
    assert out_path.stat().st_size > 0


# ── workspace group ──


@pytest.fixture
def fake_project_root(tmp_path: Path) -> Path:
    project = tmp_path / "project"
    (project / "skills" / "demo").mkdir(parents=True)
    (project / "skills" / "demo" / "SKILL.md").write_text(
        "---\nname: demo\ndescription: x\n---\nbody.\n", encoding="utf-8"
    )
    (project / "agents").mkdir(parents=True)
    (project / "agents" / "x.yaml").write_text("name: x\n", encoding="utf-8")
    return project


def test_cli_workspace_create_then_archive(tmp_path: Path, fake_project_root: Path) -> None:
    ws_root = tmp_path / "ws"
    arc_root = tmp_path / "arc"
    env = {
        "SCRIVAI_WORKSPACE_ROOT": str(ws_root),
        "SCRIVAI_ARCHIVES_ROOT": str(arc_root),
    }

    code, out, err = _run(
        [
            "workspace",
            "create",
            "--run-id",
            "cli-test",
            "--project-root",
            str(fake_project_root),
        ],
        env=env,
    )
    assert code == 0, err
    payload = json.loads(out)
    assert payload["run_id"] == "cli-test"

    # archive success
    code, out, err = _run(["workspace", "archive", "--run-id", "cli-test", "--success"], env=env)
    assert code == 0, err
    archive_path = json.loads(out)["path"]
    assert Path(archive_path).is_file()


def test_cli_workspace_missing_env_returns_error(tmp_path: Path) -> None:
    env_clean = {
        k: v
        for k, v in os.environ.items()
        if k not in ("SCRIVAI_WORKSPACE_ROOT", "SCRIVAI_ARCHIVES_ROOT")
    }
    proc = subprocess.run(
        CLI_CMD + ["workspace", "cleanup"],
        capture_output=True,
        text=True,
        env=env_clean,
    )
    assert proc.returncode == 1
    payload = _parse_error_json(proc.stderr)
    assert "SCRIVAI_WORKSPACE_ROOT" in payload["error"]


# ── trajectory group ──


@pytest.fixture
def populated_store(tmp_path: Path):
    """造一个含一次完整 run 的 trajectory db。"""
    from scrivai import TrajectoryStore

    db = tmp_path / "traj.db"
    store = TrajectoryStore(db)
    store.start_run(
        run_id="run-001",
        pes_name="extractor",
        model_name="mock-model",
        provider="mock",
        sdk_version="0.0.0",
        skills_git_hash=None,
        agents_git_hash=None,
        skills_is_dirty=False,
        task_prompt="t",
        runtime_context=None,
    )
    store.finalize_run(
        run_id="run-001",
        status="completed",
        final_output={"items": []},
        workspace_archive_path=None,
        error=None,
        error_type=None,
    )
    return db


def test_cli_trajectory_list(populated_store: Path) -> None:
    code, out, err = _run(["trajectory", "list", "--db-path", str(populated_store)])
    assert code == 0, err
    payload = json.loads(out)
    assert "runs" in payload
    assert len(payload["runs"]) == 1
    assert payload["runs"][0]["run_id"] == "run-001"


def test_cli_trajectory_get_run(populated_store: Path) -> None:
    code, out, err = _run(
        [
            "trajectory",
            "get-run",
            "--db-path",
            str(populated_store),
            "--run-id",
            "run-001",
        ]
    )
    assert code == 0, err
    payload = json.loads(out)
    assert payload["run_id"] == "run-001"
    assert payload["pes_name"] == "extractor"
    assert payload["status"] == "completed"


def test_cli_trajectory_record_feedback(tmp_path: Path, populated_store: Path) -> None:
    draft = tmp_path / "draft.json"
    draft.write_text(json.dumps({"x": 1}), encoding="utf-8")
    final = tmp_path / "final.json"
    final.write_text(json.dumps({"x": 2}), encoding="utf-8")

    code, out, err = _run(
        [
            "trajectory",
            "record-feedback",
            "--db-path",
            str(populated_store),
            "--run-id",
            "run-001",
            "--draft",
            str(draft),
            "--final",
            str(final),
            "--input-summary",
            "summary",
        ]
    )
    assert code == 0, err
    payload = json.loads(out)
    assert payload["recorded"] is True
    assert payload["run_id"] == "run-001"


def test_cli_build_eval_dataset_not_implemented(populated_store: Path) -> None:
    """build-eval-dataset 在 M2 实现,M0.75 抛 NotImplementedError。"""
    code, out, err = _run(
        [
            "trajectory",
            "build-eval-dataset",
            "--db-path",
            str(populated_store),
            "--pes-name",
            "extractor",
            "--output",
            "/tmp/out.csv",
        ]
    )
    assert code == 1
    payload = _parse_error_json(err)
    assert "M2" in payload["error"] or "not implemented" in payload["error"]
